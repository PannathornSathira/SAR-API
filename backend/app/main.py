# app/main.py

import io
import os
import csv
from datetime import datetime
from typing import List, Dict, Any, Optional
from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from langchain_core.documents import Document

import docx
import pypdf

from app.config import ALLOW_OWN_KNOWLEDGE, PORT
from app.vector_store import get_collections, create_collection, get_vector_store
from app.faq_generator import extract_faqs_from_text, expand_faq_questions, clean_rule_faqs
from app.retrieval import retrieve_and_rerank
from app.agents import rewrite_query, synthesize_answer
from app.rule_faq_generator import extract_faqs_rules
from fastapi.responses import StreamingResponse
from qdrant_client import QdrantClient
from app.config import QDRANT_URL

app = FastAPI(
    title="SAR Engine Manager & API Service",
    description="Admin UI backend and external integration gateway for the SAR FAQ system."
)

# Enable CORS for frontend API calls
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Pydantic Schemas
class CreateCollectionRequest(BaseModel):
    name: str

class FAQItem(BaseModel):
    category: str
    question: str
    answer: str
    filename: Optional[str] = ""
    source_type: Optional[str] = "LLM"

class IngestRequest(BaseModel):
    collection_name: str
    filename: str
    faqs: List[FAQItem]
    language: str = "Thai"

class QueryRequest(BaseModel):
    query: str
    chat_history: List[Dict[str, str]] = []
    allow_own_knowledge: Optional[bool] = None

@app.get("/api/v1/collections")
async def api_get_collections():
    """Lists Qdrant collections created within this app."""
    try:
        cols = get_collections()
        return {"collections": cols}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/v1/collections/stats")
async def api_get_collections_stats():
    """Returns basic stats for each valid collection in Qdrant."""
    try:
        valid_cols = get_collections()
        client = QdrantClient(url=QDRANT_URL)
        stats = []
        for col in valid_cols:
            try:
                info = client.get_collection(col)
                stats.append({
                    "name": col,
                    "points_count": info.points_count,
                    "status": info.status.value if hasattr(info.status, 'value') else str(info.status)
                })
            except Exception as e:
                print(f"[main] Error getting stats for collection {col}: {e}")
                continue
        return {"stats": stats}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/collections/create")
async def api_create_collection(req: CreateCollectionRequest):
    """Creates a new Qdrant collection configured for hybrid search."""
    if not req.name.strip():
        raise HTTPException(status_code=400, detail="Collection name cannot be empty")
    try:
        create_collection(req.name)
        return {"status": "success", "collection": req.name.strip().lower()}
    except ValueError as ve:
        raise HTTPException(status_code=400, detail=str(ve))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/extract-faq")
async def api_extract_faq(
    file: UploadFile = File(...),
    language: str = Form("Thai"),
    num_questions: int = Form(10)
):
    """Extracts clean text from a document and generates FAQ pairs."""
    filename = file.filename
    ext = filename.split(".")[-1].lower() if "." in filename else ""
    
    if ext not in ["docx", "pdf", "txt"]:
        raise HTTPException(status_code=400, detail=f"Unsupported file format: .{ext}. Supported formats are .docx, .pdf, .txt")
        
    try:
        file_bytes = await file.read()
        text = ""
        
        # 1. Clean Text Extraction
        if ext == "docx":
            doc = docx.Document(io.BytesIO(file_bytes))
            # Extract text paragraph by paragraph
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        elif ext == "pdf":
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            text_pages = []
            for i, page in enumerate(reader.pages):
                t = page.extract_text()
                if t:
                    text_pages.append(f"--- Page {i+1} ---\n{t}")
            text = "\n".join(text_pages)
        elif ext == "txt":
            text = file_bytes.decode("utf-8", errors="ignore")
            
        if not text.strip():
            raise HTTPException(status_code=400, detail="No readable text found in document.")
            
        # 2. FAQ Extraction using modular generator
        extracted_faqs_llm, extraction_cost = extract_faqs_from_text(text, filename, language, num_questions)
        extracted_faqs_rules = extract_faqs_rules(text, filename)
        
        # 3. Clean rule-based FAQs
        cleaned_rules, cleaning_cost = clean_rule_faqs(extracted_faqs_rules, language)
        total_extraction_cost = extraction_cost + cleaning_cost
        
        extracted_faqs = extracted_faqs_llm + cleaned_rules
        
        return {
            "filename": filename,
            "faqs": extracted_faqs,
            "extraction_cost": total_extraction_cost
        }
    except Exception as e:
        print(f"[main] Error during text extraction/FAQ generation: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/collections/ingest")
async def api_ingest_faqs(req: IngestRequest):
    """Embeds and saves the FAQ pairs to the chosen Qdrant collection."""
    if not req.faqs:
        raise HTTPException(status_code=400, detail="FAQ list is empty")
        
    try:
        # 1. Expand FAQs
        faqs_to_expand = []
        for faq in req.faqs:
            d = faq.model_dump() if hasattr(faq, 'model_dump') else faq.dict()
            d['original_question'] = faq.question
            faqs_to_expand.append(d)
            
        expanded_faqs, expansion_cost = expand_faq_questions(faqs_to_expand, language=req.language)

        # 2. Create documents
        docs = []
        for faq in expanded_faqs:
            page_content = f"Question: {faq['question']}\nAnswer: {faq['answer']}"
            metadata = {
                "category": faq.get('category', ''),
                "original_question": faq.get('original_question', faq.get('question', '')),
                "answer": faq.get('answer', ''),
                "source_file": faq.get('filename') or req.filename,
                "source_type": faq.get('source_type', 'LLM')
            }
            docs.append(Document(page_content=page_content, metadata=metadata))
            
        # Add to vector store in hybrid mode
        store = get_vector_store(req.collection_name)
        store.add_documents(docs)
        
        # 3. Export to CSV
        os.makedirs("exports", exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        export_filename = f"exports/{req.collection_name}_{timestamp}.csv"
        
        with open(export_filename, mode='w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(["Category", "Filename", "Original Question", "Question Variation", "Answer", "Source Type"])
            for faq in expanded_faqs:
                writer.writerow([
                    faq.get('category', ''),
                    faq.get('filename', req.filename),
                    faq.get('original_question', ''),
                    faq.get('question', ''),
                    faq.get('answer', ''),
                    faq.get('source_type', 'LLM')
                ])
        
        return {"status": "success", "count": len(docs), "expansion_cost": expansion_cost, "csv_path": export_filename}
    except Exception as e:
        print(f"[main] Error during ingestion: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/query/{collection_name}")
async def api_query_collection(collection_name: str, req: QueryRequest):
    """Exposes a dynamic SAR query endpoint for external application queries."""
    # Ensure collection exists
    collection_name = collection_name.strip().lower()
    valid_cols = get_collections()
    if collection_name not in valid_cols:
        raise HTTPException(status_code=404, detail=f"Collection '{collection_name}' not found or not initialized in this app")
        
    try:
        effective_allow_own_knowledge = (
            ALLOW_OWN_KNOWLEDGE
            if req.allow_own_knowledge is None
            else req.allow_own_knowledge
        )

        # 1. Query Rewrite (using memory)
        rewritten = rewrite_query(req.query, req.chat_history)
        
        # 2. Retrieve & Rerank (using tokenization and CrossEncoder)
        contexts, best_score = retrieve_and_rerank(collection_name, rewritten)
        
        # 3. Write Answer (Threshold logic + Fallback)
        response_text, match_type = synthesize_answer(
            rewritten,
            contexts,
            best_score,
            allow_own_knowledge=effective_allow_own_knowledge
        )
        
        # Format sources
        sources = list(set([c["source_file"] for c in contexts]))
        
        # Format context for return JSON
        formatted_context = [
            {
                "content": c["content"],
                "source_file": c["source_file"],
                "category": c["category"]
            }
            for c in contexts
        ]
        
        return {
            "response": response_text,
            "match_type": match_type,
            "score": best_score,
            "allow_own_knowledge": effective_allow_own_knowledge,
            "sources": sources,
            "context": formatted_context
        }
    except Exception as e:
        print(f"[main] Error querying collection '{collection_name}': {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/v1/collections/export/{collection_name}")
async def api_export_collection(collection_name: str):
    """Exports all points in a collection to a CSV file."""
    try:
        client = QdrantClient(url=QDRANT_URL)
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(["Category", "Filename", "Original Question", "Question", "Answer", "Source Type"])
        
        offset = None
        while True:
            records, next_page = client.scroll(
                collection_name=collection_name,
                limit=100,
                offset=offset,
                with_payload=True
            )
            
            for record in records:
                payload = record.payload
                metadata = payload.get("metadata", {})
                
                cat = metadata.get("category", "")
                fname = metadata.get("source_file", "")
                orig_q = metadata.get("original_question", "")
                ans = metadata.get("answer", "")
                stype = metadata.get("source_type", "LLM")
                
                page_content = payload.get("page_content", "")
                lines = page_content.split("\n")
                q_var = ""
                for line in lines:
                    if line.startswith("Question: "):
                        q_var = line[len("Question: "):]
                        break
                        
                writer.writerow([cat, fname, orig_q, q_var, ans, stype])
                
            if next_page is None:
                break
            offset = next_page
            
        output.seek(0)
        return StreamingResponse(
            iter([output.getvalue()]), 
            media_type="text/csv", 
            headers={"Content-Disposition": f"attachment; filename={collection_name}_export.csv"}
        )
    except Exception as e:
        print(f"[main] Error exporting collection '{collection_name}': {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/v1/collections/{collection_name}/faqs")
async def api_get_collection_faqs(collection_name: str, limit: int = 100):
    """Fetches a sample of FAQs from a collection."""
    try:
        client = QdrantClient(url=QDRANT_URL)
        records, _ = client.scroll(
            collection_name=collection_name,
            limit=limit,
            with_payload=True
        )
        
        faqs = []
        for record in records:
            payload = record.payload
            metadata = payload.get("metadata", {})
            page_content = payload.get("page_content", "")
            
            q_var = ""
            lines = page_content.split("\n")
            for line in lines:
                if line.startswith("Question: "):
                    q_var = line[len("Question: "):]
                    break
                    
            faqs.append({
                "id": str(record.id),
                "category": metadata.get("category", ""),
                "original_question": metadata.get("original_question", ""),
                "question": q_var,
                "answer": metadata.get("answer", ""),
                "filename": metadata.get("source_file", ""),
                "source_type": metadata.get("source_type", "")
            })
            
        return {"faqs": faqs}
    except Exception as e:
        print(f"[main] Error getting FAQs for '{collection_name}': {e}")
        raise HTTPException(status_code=500, detail=str(e))
